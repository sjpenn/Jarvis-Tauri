"""
Document Processor - Handle various document formats for training

Supports:
- PDF files
- Text/Markdown files
- HTML content
- Chunking with overlap for better context
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional
import hashlib


@dataclass
class DocumentChunk:
    """A chunk of a document"""
    content: str
    start_pos: int
    end_pos: int
    metadata: dict = field(default_factory=dict)


@dataclass
class Document:
    """A processed document"""
    id: str
    source_path: Path
    content: str
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    
    @staticmethod
    def generate_id(source_path: Path) -> str:
        """Generate a unique ID for a document"""
        return hashlib.md5(str(source_path).encode()).hexdigest()


class DocumentProcessor:
    """Process documents into chunks for Q&A generation"""
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 128,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, file_path: Path) -> Document:
        """
        Process a document file into chunks.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Processed Document
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Determine file type and extract content
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            content = self._extract_pdf(file_path)
        elif suffix in ['.txt', '.md', '.markdown']:
            content = self._extract_text(file_path)
        elif suffix in ['.html', '.htm']:
            content = self._extract_html(file_path)
        elif suffix in ['.docx', '.doc']:
            content = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        # Create document
        doc = Document(
            id=Document.generate_id(file_path),
            source_path=file_path,
            content=content,
            metadata={
                "filename": file_path.name,
                "file_type": suffix,
                "size_bytes": file_path.stat().st_size,
            }
        )
        
        # Chunk the content
        doc.chunks = self._chunk_text(content)
        
        return doc
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from a text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file"""
        try:
            import pypdf
            
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return "\n\n".join(text)
        except ImportError:
            raise ImportError(
                "PDF support requires pypdf. Install with: pip install pypdf"
            )
    
    def _extract_html(self, file_path: Path) -> str:
        """Extract text from an HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
        except ImportError:
            raise ImportError(
                "HTML support requires beautifulsoup4. Install with: pip install beautifulsoup4"
            )
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from a DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text.append(row_text)
            
            return '\n\n'.join(text)
        except ImportError:
            raise ImportError(
                "DOCX support requires python-docx. Install with: pip install python-docx"
            )
    
    def _chunk_text(self, text: str) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        # Split into sentences (simple approach)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        current_chunk = []
        current_length = 0
        start_pos = 0
        
        for sentence in sentences:
            sentence_length = len(sentence.split())
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    start_pos=start_pos,
                    end_pos=start_pos + len(chunk_text),
                ))
                
                # Start new chunk with overlap
                overlap_sentences = []
                overlap_length = 0
                for s in reversed(current_chunk):
                    if overlap_length + len(s.split()) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_length += len(s.split())
                    else:
                        break
                
                start_pos += len(chunk_text) - len(" ".join(overlap_sentences))
                current_chunk = overlap_sentences
                current_length = overlap_length
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Add final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(DocumentChunk(
                content=chunk_text,
                start_pos=start_pos,
                end_pos=start_pos + len(chunk_text),
            ))
        
        return chunks
